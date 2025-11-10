"""DOC/DOCX file chunker implementation."""

import io
from typing import List
from logging import Logger, getLogger

from .base import BaseChunker, ChunkResult


class DocChunker(BaseChunker):
    """Chunker for Microsoft Word documents (.doc, .docx)."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize DOC chunker.
        
        Args:
            chunk_size: Maximum size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
        """
        super().__init__(chunk_size, chunk_overlap)
        self._logger: Logger = getLogger(__name__)
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOC/DOCX file.
        
        Args:
            file_content: Raw file bytes
            filename: Name of the file
            
        Returns:
            Extracted text content
        """
        try:
            # Import python-docx for .docx files
            from docx import Document
            
            # Create document from bytes
            doc_stream = io.BytesIO(file_content)
            doc = Document(doc_stream)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            self._logger.info(f"Extracted {len(full_text)} characters from {filename}")
            return full_text
            
        except ImportError:
            self._logger.error("python-docx library not installed. Install it with: pip install python-docx")
            raise ValueError("python-docx library is required for DOC/DOCX processing")
        except Exception as e:
            self._logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    async def chunk_text(self, text: str) -> List[ChunkResult]:
        """Split text into overlapping chunks.
        
        Args:
            text: Text to split into chunks
            
        Returns:
            List of ChunkResult objects
        """
        if not text:
            return []
        
        chunks = []
        chunk_index = 0
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If this is not the last chunk, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence boundaries (. ! ?)
                sentence_end = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end)
                )
                
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundary (space)
                    space_pos = text.rfind(' ', start, end)
                    if space_pos > start:
                        end = space_pos
            
            # Extract chunk
            chunk_content = text[start:end].strip()
            
            if chunk_content:
                chunk_result = ChunkResult(
                    content=chunk_content,
                    chunk_index=chunk_index,
                    metadata={
                        'start_char': start,
                        'end_char': end,
                        'length': len(chunk_content)
                    }
                )
                chunks.append(chunk_result)
                chunk_index += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap if end < len(text) else end
            
            # Ensure we make progress
            if start <= end - self.chunk_size:
                start = end
        
        self._logger.info(f"Created {len(chunks)} chunks from text of length {len(text)}")
        return chunks
